from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── Colour palette ──
NAVY   = RGBColor(0x0D, 0x11, 0x17)   # near-black
BLUE   = RGBColor(0x38, 0x8B, 0xFD)   # accent blue
GREEN  = RGBColor(0x3F, 0xB9, 0x50)   # accent green
GREY   = RGBColor(0x8B, 0x94, 0x9E)   # secondary text
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF6, 0xF8, 0xFA)   # light background
DARK   = RGBColor(0x16, 0x1B, 0x22)   # surface

# ── Helpers ──
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        val = kwargs.get(side, {})
        if val:
            el = OxmlElement(f'w:{side}')
            for k,v in val.items():
                el.set(qn(f'w:{k}'), v)
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, bold=True, size=16, color=NAVY, font='Calibri')
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), '388BFD')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, text, bold=True, size=13, color=RGBColor(0x1F,0x2D,0x3D), font='Calibri')
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    add_run(p, text, size=11, color=RGBColor(0x24,0x29,0x2E))
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.6 + level*0.5)
    add_run(p, text, size=11, color=RGBColor(0x24,0x29,0x2E))
    return p

def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'),   'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'),  '1C2128')
        p._p.get_or_add_pPr().append(shading)
        run = p.add_run(line if line else ' ')
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0xE6,0xED,0xF3)

def info_box(label, value):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(11)
    r = tbl.rows[0]
    r.cells[0].text = label
    r.cells[1].text = value
    set_cell_bg(r.cells[0], '0D1117')
    set_cell_bg(r.cells[1], 'F6F8FA')
    for run in r.cells[0].paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(10)
    for run in r.cells[1].paragraphs[0].runs:
        run.font.size  = Pt(10)
    doc.add_paragraph()

# ════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
add_run(cover, 'CERTISHIELD', bold=True, size=36, color=BLUE, font='Calibri')

cover2 = doc.add_paragraph()
cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(cover2, 'Blockchain-Based Certificate Verification System', bold=False,
        size=16, color=GREY, font='Calibri')

doc.add_paragraph()
cover3 = doc.add_paragraph()
cover3.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(cover3, 'Project Report', bold=True, size=14, color=NAVY, font='Calibri')

doc.add_paragraph()
meta = [
    ('Author',           'Dhruv Rajpal'),
    ('Live Demo',        'https://certishield.vercel.app'),
    ('Network',          'Ethereum Sepolia Testnet'),
    ('Contract Address', '0xc16Fc9Ea835930CC4042b85D1c73ED82D3761A88'),
    ('Technology Stack', 'Solidity · Web3.js · Flask · Python · HTML/CSS/JS'),
    ('Date',             'May 2026'),
]
for label, value in meta:
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(mp, f'{label}: ', bold=True, size=11, color=NAVY)
    add_run(mp, value, size=11, color=GREY)

doc.add_page_break()

# ════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════
toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(toc_title, 'Table of Contents', bold=True, size=18, color=NAVY)
doc.add_paragraph()

toc_items = [
    ('1.', 'Project Summary'),
    ('2.', 'Problem Statement'),
    ('3.', 'Objectives'),
    ('4.', 'Scope'),
    ('5.', 'Tools & Environment'),
    ('6.', 'System Design'),
    ('7.', 'Implementation'),
    ('8.', 'Code'),
    ('9.', 'Running the Program'),
    ('10.','Sample Output'),
    ('11.','Testing & Validation'),
    ('12.','Ethical & Legal Considerations'),
    ('13.','Conclusion'),
    ('14.','References'),
]
for num, title in toc_items:
    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(4)
    add_run(tp, f'  {num}  ', bold=True, size=11, color=BLUE)
    add_run(tp, title, size=11, color=NAVY)

doc.add_page_break()

# ════════════════════════════════════════
# 1. PROJECT SUMMARY
# ════════════════════════════════════════
heading1('1. Project Summary')
body(
    'CertiShield is a decentralised Web3 application designed to eliminate certificate fraud '
    'by storing cryptographic hashes of institutional credentials on the Ethereum Sepolia '
    'blockchain. Rather than storing any document on-chain, CertiShield computes a SHA-256 '
    'hash of a PDF in the user\'s browser and records only that fingerprint through a Solidity '
    'smart contract. The same process is repeated for verification—if the hash matches the '
    'on-chain record, the certificate is confirmed authentic.'
)
body(
    'The project is live at https://certishield.vercel.app and consists of a Python/Flask '
    'backend serving a modern, professional single-page frontend that communicates with '
    'Ethereum via MetaMask and Web3.js.'
)

tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = ['Attribute', 'Detail']
rows_data = [
    ('Project Name',    'CertiShield'),
    ('Type',            'Decentralised Application (dApp)'),
    ('Blockchain',      'Ethereum Sepolia Testnet'),
    ('Live URL',        'https://certishield.vercel.app'),
    ('Smart Contract',  '0xc16Fc9Ea835930CC4042b85D1c73ED82D3761A88'),
]
header_row = tbl.rows[0]
set_cell_bg(header_row.cells[0], '0D1117')
set_cell_bg(header_row.cells[1], '0D1117')
for i, h in enumerate(headers):
    header_row.cells[i].text = h
    for run in header_row.cells[i].paragraphs[0].runs:
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)

for idx, (k, v) in enumerate(rows_data):
    row = tbl.rows[idx+1] if idx+1 < len(tbl.rows) else tbl.add_row()
    # rebuild if needed
    row = tbl.rows[idx+1]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], 'F6F8FA')
    for run in row.cells[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(10)
    for run in row.cells[1].paragraphs[0].runs:
        run.font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ════════════════════════════════════════
heading1('2. Problem Statement')
body(
    'Academic and professional certificates are vulnerable to forgery and tampering. '
    'Traditional verification methods—phone calls, physical stamps, and centralised '
    'databases—are slow, expensive, and prone to single points of failure. A forged '
    'certificate produced with modern image-editing tools is difficult to distinguish '
    'from a genuine one without contacting the issuing authority.'
)
body(
    'Centralised verification systems face additional risks:'
)
for item in [
    'Data breaches exposing sensitive credential information',
    'Server downtime rendering verification impossible',
    'Lack of transparency—verifiers must trust a third party',
    'High operational cost for institutions maintaining verification infrastructure',
]:
    bullet(item)

body(
    'There is a clear need for a tamper-proof, publicly auditable, and low-cost '
    'verification mechanism that removes the reliance on any single institution or server.'
)

# ════════════════════════════════════════
# 3. OBJECTIVES
# ════════════════════════════════════════
heading1('3. Objectives')
objectives = [
    ('Primary Objective',
     'Build a decentralised certificate verification system that records SHA-256 hashes '
     'of PDF certificates on the Ethereum blockchain, making forgery cryptographically impossible.'),
    ('Transparency',
     'Ensure every registration and verification event is publicly auditable on Etherscan '
     'without requiring trust in any centralised authority.'),
    ('Privacy',
     'Hash certificates locally in the browser so that the original document is never '
     'transmitted to any server.'),
    ('Accessibility',
     'Provide an intuitive, professional UI that any institution or individual can use '
     'without blockchain expertise beyond a MetaMask wallet.'),
    ('Cost Efficiency',
     'Leverage an Ethereum testnet (Sepolia) during development and keep on-chain data '
     'minimal (string hash only) to minimise gas costs in production.'),
]
for title, desc in objectives:
    heading2(f'  {title}')
    body(desc, indent=True)

# ════════════════════════════════════════
# 4. SCOPE
# ════════════════════════════════════════
heading1('4. Scope')
heading2('In Scope')
for item in [
    'PDF certificate hashing (SHA-256) entirely client-side',
    'On-chain registration of certificate hashes via Solidity smart contract',
    'Public verification of certificate authenticity against on-chain data',
    'MetaMask wallet integration for transaction signing',
    'Professional single-page web application served via Flask/Vercel',
    'Transaction confirmation with Etherscan deep-link',
]:
    bullet(item)

heading2('Out of Scope')
for item in [
    'Storage of actual certificate files (on-chain or off-chain)',
    'Multi-signature or role-based access for large institutions (future work)',
    'Mainnet Ethereum deployment (requires ETH—testnet only in this version)',
    'Mobile native application',
    'Certificate issuance workflow (only verification is handled)',
]:
    bullet(item)

# ════════════════════════════════════════
# 5. TOOLS & ENVIRONMENT
# ════════════════════════════════════════
heading1('5. Tools & Environment')

tools = [
    ('Category',          'Tool / Technology',         'Version / Detail'),
    ('Smart Contract',    'Solidity',                  '^0.8.0'),
    ('Blockchain',        'Ethereum Sepolia Testnet',  'ChainID 11155111'),
    ('Wallet',            'MetaMask Browser Extension','Latest'),
    ('Web3 Library',      'Web3.js',                   'CDN (latest)'),
    ('Backend',           'Python / Flask',            'Python 3.12 / Flask 3.x'),
    ('Production Server', 'Gunicorn',                  '25.x'),
    ('Frontend',          'HTML5 / CSS3 / Vanilla JS', 'ES2022'),
    ('Fonts',             'Inter & JetBrains Mono',    'Google Fonts CDN'),
    ('Hashing',           'Web Crypto API (SHA-256)',  'Browser-native'),
    ('IDE / Dev',         'Replit',                    'Cloud-based'),
    ('Hosting',           'Vercel',                    'Free tier'),
    ('Version Control',   'Git / GitHub',              'github.com/rustyprophet213'),
    ('Contract Explorer', 'Etherscan Sepolia',         'sepolia.etherscan.io'),
]

t = doc.add_table(rows=len(tools), cols=3)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.LEFT
for r_idx, row_data in enumerate(tools):
    row = t.rows[r_idx]
    for c_idx, cell_text in enumerate(row_data):
        row.cells[c_idx].text = cell_text
        if r_idx == 0:
            set_cell_bg(row.cells[c_idx], '0D1117')
            for run in row.cells[c_idx].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
        else:
            bg = 'F6F8FA' if r_idx % 2 == 0 else 'FFFFFF'
            set_cell_bg(row.cells[c_idx], bg)
            for run in row.cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(10)
                if c_idx == 0:
                    run.font.bold = True

doc.add_paragraph()

# ════════════════════════════════════════
# 6. SYSTEM DESIGN
# ════════════════════════════════════════
heading1('6. System Design')

heading2('6.1 Architecture Overview')
body(
    'CertiShield follows a three-tier architecture: a thin Flask web server, a '
    'client-side JavaScript frontend, and an Ethereum smart contract acting as the '
    'immutable data layer.'
)

arch = [
    ('Tier',            'Component',         'Responsibility'),
    ('Presentation',    'HTML/CSS/JS (SPA)', 'User interface, file hashing, MetaMask interaction'),
    ('Application',     'Flask (Python)',     'Serves static assets and the HTML template'),
    ('Data / Logic',    'Solidity Contract', 'Stores and queries certificate hashes on-chain'),
]
ta = doc.add_table(rows=len(arch), cols=3)
ta.style = 'Table Grid'
for r_idx, row_data in enumerate(arch):
    row = ta.rows[r_idx]
    for c_idx, txt in enumerate(row_data):
        row.cells[c_idx].text = txt
        if r_idx == 0:
            set_cell_bg(row.cells[c_idx], '161B22')
            for run in row.cells[c_idx].paragraphs[0].runs:
                run.font.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(10)
        else:
            for run in row.cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(10)
doc.add_paragraph()

heading2('6.2 Smart Contract Design')
body(
    'The Solidity contract CertificateVerification maintains a private mapping from '
    'string (SHA-256 hash) to bool. Only the deploying admin address can register '
    'hashes; anyone can call verifyCertificate() as a read-only (view) function at '
    'no gas cost.'
)

heading2('6.3 Client-Side Hashing Flow')
for step in [
    '1. User selects a PDF file via drag-and-drop or file picker.',
    '2. JavaScript reads the file as an ArrayBuffer.',
    '3. The Web Crypto API computes a SHA-256 digest of the buffer in-browser.',
    '4. The 64-character hex string is passed to the smart contract via Web3.js.',
    '5. MetaMask prompts the user to sign and submit the transaction.',
    '6. On confirmation, the transaction hash is displayed with an Etherscan link.',
]:
    bullet(step)

heading2('6.4 Data Flow Diagram (Text Representation)')
body('Registration:')
code_block([
    '  [User PDF] ──SHA-256──► [hex hash] ──Web3.js──► [MetaMask] ──Tx──► [Ethereum]',
    '                                                                          │',
    '                                              Smart Contract stores hash ◄┘',
])
doc.add_paragraph()
body('Verification:')
code_block([
    '  [User PDF] ──SHA-256──► [hex hash] ──Web3.js──► [Contract.verifyCertificate()]',
    '                                                          │',
    '                                              Returns true / false ──► [UI Result]',
])
doc.add_paragraph()

# ════════════════════════════════════════
# 7. IMPLEMENTATION
# ════════════════════════════════════════
heading1('7. Implementation')

heading2('7.1 Smart Contract')
body(
    'The contract is written in Solidity 0.8.0 and deployed to the Sepolia testnet. '
    'It uses a private mapping to prevent direct enumeration of stored hashes, and '
    'an onlyAdmin modifier to restrict registration to the deploying address.'
)

heading2('7.2 Flask Backend')
body(
    'The Flask application has a single route (/) that renders the Jinja2 HTML template. '
    'All blockchain logic runs client-side; Flask\'s only responsibility is serving '
    'the page and static assets (CSS, JS). Gunicorn is used as the WSGI server in '
    'production, bound to 0.0.0.0:5000.'
)

heading2('7.3 Frontend UI')
body(
    'The frontend is a single-page application built with vanilla HTML, CSS, and JavaScript. '
    'Key UI/UX features implemented:'
)
for item in [
    'Sticky navbar with live network status indicator (animated green dot)',
    'Drag-and-drop PDF upload zones with hover states and file-name confirmation',
    'Toast notification system replacing browser alert() calls',
    'Full-screen loading overlay with spinner during blockchain transactions',
    'Transaction hash display with one-click clipboard copy',
    'Etherscan deep-link for every confirmed transaction',
    'Colour-coded result card (green = authentic, red = invalid)',
    '"How It Works" four-step explainer at the bottom of the page',
]:
    bullet(item)

heading2('7.4 Design System')
body(
    'The UI uses a GitHub-inspired dark design system with CSS custom properties '
    '(variables) for colours, spacing, and shadows. Typography uses Inter for body '
    'text and JetBrains Mono for code/hash display.'
)

# ════════════════════════════════════════
# 8. CODE
# ════════════════════════════════════════
heading1('8. Code')

heading2('8.1 Smart Contract — CertificateVerification.sol')
code_block([
    '// SPDX-License-Identifier: MIT',
    'pragma solidity ^0.8.0;',
    '',
    'contract CertificateVerification {',
    '',
    '    mapping(string => bool) private certificates;',
    '    address public admin;',
    '',
    '    constructor() {',
    '        admin = msg.sender;',
    '    }',
    '',
    '    modifier onlyAdmin() {',
    '        require(msg.sender == admin, "Not authorized");',
    '        _;',
    '    }',
    '',
    '    function addCertificate(string memory hash) public onlyAdmin {',
    '        certificates[hash] = true;',
    '    }',
    '',
    '    function verifyCertificate(string memory hash) public view returns (bool) {',
    '        return certificates[hash];',
    '    }',
    '}',
])
doc.add_paragraph()

heading2('8.2 Flask Application — app.py')
code_block([
    'from flask import Flask, render_template',
    '',
    'app = Flask(__name__)',
    '',
    '@app.route("/")',
    'def home():',
    '    return render_template("index.html")',
    '',
    'if __name__ == "__main__":',
    '    app.run(host="0.0.0.0", port=5000, debug=True)',
])
doc.add_paragraph()

heading2('8.3 Key JavaScript Functions — app.js (excerpts)')
code_block([
    '// SHA-256 hashing (browser-native, no library needed)',
    'async function generateHash(file) {',
    '    const buffer = await file.arrayBuffer();',
    '    const hashBuffer = await crypto.subtle.digest("SHA-256", buffer);',
    '    return Array.from(new Uint8Array(hashBuffer))',
    '        .map(b => b.toString(16).padStart(2, "0")).join("");',
    '}',
    '',
    '// Register certificate on blockchain',
    'async function addCertificate() {',
    '    const file = document.getElementById("fileInput").files[0];',
    '    const hash = await generateHash(file);',
    '    showLoader(true, "Hashing & sending transaction...");',
    '    const tx = await contract.methods',
    '        .addCertificate(hash).send({ from: account });',
    '    showLoader(false);',
    '    // display tx hash + Etherscan link ...',
    '}',
    '',
    '// Verify certificate authenticity',
    'async function verifyCertificate() {',
    '    const file = document.getElementById("verifyFileInput").files[0];',
    '    const hash = await generateHash(file);',
    '    const isValid = await contract.methods',
    '        .verifyCertificate(hash).call();',
    '    // show green (authentic) or red (invalid) result card ...',
    '}',
])
doc.add_paragraph()

# ════════════════════════════════════════
# 9. RUNNING THE PROGRAM
# ════════════════════════════════════════
heading1('9. Running the Program')

heading2('9.1 Prerequisites')
for item in [
    'Python 3.10+ installed',
    'MetaMask browser extension installed and configured for Sepolia Testnet',
    'Sepolia test ETH in your wallet (available from https://sepoliafaucet.com)',
    'Git installed (optional, for cloning)',
]:
    bullet(item)

heading2('9.2 Local Setup')
code_block([
    '# 1. Clone the repository',
    'git clone https://github.com/rustyprophet213/Certishield.git',
    'cd Certishield',
    '',
    '# 2. Install dependencies',
    'pip install -r requirements.txt',
    '',
    '# 3. Run the development server',
    'python app.py',
    '',
    '# App is now available at http://localhost:5000',
])
doc.add_paragraph()

heading2('9.3 Production (Gunicorn)')
code_block([
    'gunicorn --bind 0.0.0.0:5000 --reuse-port app:app',
])
doc.add_paragraph()

heading2('9.4 Usage Walkthrough')
steps = [
    ('Step 1 — Connect Wallet',
     'Click "Connect MetaMask". Approve the connection request in the MetaMask popup. '
     'Your shortened wallet address will appear in the navbar badge.'),
    ('Step 2 — Register a Certificate',
     'Drag a PDF onto the Register drop zone (or click Browse). Click "Register on Blockchain". '
     'Confirm the transaction in MetaMask. The transaction hash and an Etherscan link appear on success.'),
    ('Step 3 — Verify a Certificate',
     'Drag the same or a different PDF onto the Verify drop zone. Click "Verify Authenticity". '
     'A green card confirms the certificate is authentic; a red card indicates no on-chain record exists.'),
]
for title, desc in steps:
    heading2(f'  {title}')
    body(desc, indent=True)

# ════════════════════════════════════════
# 10. SAMPLE OUTPUT
# ════════════════════════════════════════
heading1('10. Sample Output')

heading2('10.1 Successful Registration')
code_block([
    'Transaction Hash:',
    '0x4a7f3c8b2e1d9f0a6c5b3e2d1f8a7c4b9e0f3a2d1c6b5e4f3a2b1c0d9e8f7a6',
    '',
    'Etherscan Link:',
    'https://sepolia.etherscan.io/tx/0x4a7f3c8b...',
    '',
    'Toast notification: "Certificate registered on the blockchain!"  ✅',
])
doc.add_paragraph()

heading2('10.2 Authentic Certificate Verification')
code_block([
    'Result Card (green):',
    '  ✅  Authentic Certificate',
    '       This certificate hash is verified on-chain.',
    '',
    'Toast notification: "Certificate verified as authentic."  ✅',
])
doc.add_paragraph()

heading2('10.3 Invalid Certificate Verification')
code_block([
    'Result Card (red):',
    '  ❌  Invalid Certificate',
    '       No matching record found on the blockchain.',
    '',
    'Toast notification: "Certificate could not be verified."  ❌',
])
doc.add_paragraph()

heading2('10.4 No Wallet Connected')
code_block([
    'Toast notification: "Please connect your wallet first."  ❌',
])
doc.add_paragraph()

# ════════════════════════════════════════
# 11. TESTING & VALIDATION
# ════════════════════════════════════════
heading1('11. Testing & Validation')

test_cases = [
    ('Test Case', 'Input', 'Expected Output', 'Result'),
    ('TC-01: Register valid PDF',
     'Legitimate certificate PDF + connected wallet',
     'Transaction confirmed, hash stored on-chain',
     'PASS'),
    ('TC-02: Verify registered PDF',
     'Same PDF used in TC-01',
     'Green "Authentic Certificate" result card',
     'PASS'),
    ('TC-03: Verify unregistered PDF',
     'Different PDF not previously registered',
     'Red "Invalid Certificate" result card',
     'PASS'),
    ('TC-04: No wallet connected',
     'Click Register without MetaMask',
     'Toast error "Please connect your wallet first"',
     'PASS'),
    ('TC-05: No file selected',
     'Click Register with no PDF chosen',
     'Toast error "Please select a PDF certificate"',
     'PASS'),
    ('TC-06: Non-PDF file dropped',
     'Drag a .jpg onto upload zone',
     'Toast error "Only PDF files are accepted"',
     'PASS'),
    ('TC-07: Transaction rejected',
     'User rejects MetaMask prompt',
     'Toast error "Transaction failed or was rejected"',
     'PASS'),
    ('TC-08: Hash determinism',
     'Same PDF registered twice',
     'Second tx succeeds; verify returns true both times',
     'PASS'),
    ('TC-09: Tampered PDF',
     'PDF with one byte changed',
     'Different hash → verify returns Invalid',
     'PASS'),
    ('TC-10: Drag-and-drop',
     'PDF dragged directly onto drop zone',
     'File name shown, upload zone highlights on dragover',
     'PASS'),
]
tt = doc.add_table(rows=len(test_cases), cols=4)
tt.style = 'Table Grid'
tt.alignment = WD_TABLE_ALIGNMENT.LEFT
for r_idx, row_data in enumerate(test_cases):
    row = tt.rows[r_idx]
    for c_idx, txt in enumerate(row_data):
        row.cells[c_idx].text = txt
        if r_idx == 0:
            set_cell_bg(row.cells[c_idx], '0D1117')
            for run in row.cells[c_idx].paragraphs[0].runs:
                run.font.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(9)
        else:
            bg = 'F6F8FA' if r_idx % 2 == 0 else 'FFFFFF'
            set_cell_bg(row.cells[c_idx], bg)
            for run in row.cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(9)
                if txt == 'PASS':
                    run.font.color.rgb = GREEN
                    run.font.bold = True
doc.add_paragraph()

# ════════════════════════════════════════
# 12. ETHICAL & LEGAL CONSIDERATIONS
# ════════════════════════════════════════
heading1('12. Ethical & Legal Considerations')

heading2('12.1 Privacy')
body(
    'CertiShield is designed with privacy as a first principle. Certificate files are '
    'hashed entirely within the user\'s browser using the Web Crypto API. No document '
    'content is ever transmitted to the Flask server or stored anywhere off-chain. '
    'Only a 64-character hex string reaches the blockchain, which cannot be reversed '
    'to reconstruct the original document.'
)

heading2('12.2 Data Immutability')
body(
    'Once a hash is written to the Ethereum blockchain, it cannot be deleted or altered. '
    'Institutions deploying this system should have clear policies on which certificates '
    'are eligible for registration, as revocation is not natively supported in this version. '
    'Future iterations should implement a revocation mapping.'
)

heading2('12.3 Admin Key Security')
body(
    'The smart contract restricts certificate registration to the admin address (the '
    'deployer). If the admin\'s private key is compromised, unauthorised hashes could '
    'be added. In a production system, a multi-signature wallet (e.g., Gnosis Safe) '
    'should hold the admin role.'
)

heading2('12.4 Testnet vs Mainnet')
body(
    'This project is deployed on the Sepolia testnet. Testnet transactions have no '
    'real monetary value and the network may be reset. A production deployment requires '
    'mainnet ETH for gas, real-world contract auditing, and compliance with local '
    'regulations around digital record-keeping (e.g., GDPR, IT Act 2000 in India).'
)

heading2('12.5 Accessibility')
body(
    'The UI uses semantic HTML, sufficient colour contrast ratios, and keyboard-accessible '
    'interactive elements to meet basic WCAG 2.1 Level AA guidelines.'
)

# ════════════════════════════════════════
# 13. CONCLUSION
# ════════════════════════════════════════
heading1('13. Conclusion')
body(
    'CertiShield demonstrates that blockchain technology can provide a practical, '
    'low-cost, and tamper-proof solution to the widespread problem of certificate '
    'fraud. By combining SHA-256 browser-side hashing with an immutable Ethereum '
    'smart contract, the system guarantees that no registered certificate can be '
    'silently altered without invalidating its on-chain record.'
)
body(
    'The project successfully achieves all stated objectives: privacy-preserving '
    'hashing, transparent on-chain storage, public verification, and a professional '
    'user interface accessible to non-technical users. All ten test cases passed, '
    'confirming correct behaviour across normal and edge-case scenarios.'
)
body(
    'Future enhancements could include: certificate revocation, role-based institutional '
    'access, IPFS integration for optional document storage, and mainnet deployment '
    'with a formal smart contract audit.'
)

# ════════════════════════════════════════
# 14. REFERENCES
# ════════════════════════════════════════
heading1('14. References')
refs = [
    '[1]  Ethereum Foundation. "Solidity Documentation." https://docs.soliditylang.org/',
    '[2]  Web3.js. "Web3.js Documentation." https://web3js.readthedocs.io/',
    '[3]  MetaMask. "MetaMask Developer Documentation." https://docs.metamask.io/',
    '[4]  NIST. "Secure Hash Standard (SHA-2)." FIPS PUB 180-4, 2015.',
    '[5]  Mozilla Developer Network. "Web Crypto API." https://developer.mozilla.org/en-US/docs/Web/API/Web_Crypto_API',
    '[6]  Etherscan Sepolia. "Sepolia Testnet Explorer." https://sepolia.etherscan.io/',
    '[7]  Vercel. "Vercel Deployment Documentation." https://vercel.com/docs',
    '[8]  Pallets Projects. "Flask Documentation." https://flask.palletsprojects.com/',
    '[9]  Gunicorn. "Gunicorn WSGI HTTP Server." https://gunicorn.org/',
    '[10] GitHub Repository. "CertiShield Source Code." https://github.com/rustyprophet213/Certishield',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run(p, ref, size=10, color=RGBColor(0x24,0x29,0x2E))

# ── Final credit ──
doc.add_paragraph()
credit = doc.add_paragraph()
credit.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(credit, 'Made with ❤️ by Dhruv Rajpal  •  CertiShield  •  2026',
        size=10, color=GREY, italic=True)

# ── Save ──
path = 'CertiShield_Project_Report.docx'
doc.save(path)
print(f'Report saved: {path}')
